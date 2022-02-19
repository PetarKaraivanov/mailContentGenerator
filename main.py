""" modifying a docx template using a csv to
    replace placeholders inside the docx and to create
    a pdf.
"""
import os
from docx import Document
from docx2pdf import convert

def get_info():
    """ shops.csv holds the information, that will be replaced.
        First line holds the keys.
        Extract the keys and create a list of dictionaries.
        Structure of the dictionary:
            key: the word, which will be replaced in the docx template.
            value: the real information which will be inserted into the docx.
    """
    lst_csv_data = []
    with open("shops.csv", "r", encoding="utf-8") as shops_file:
        lst_shops = shops_file.readlines()
        #get the keys
        lst_keys = lst_shops.pop(0).replace("\n", "").split(";")
        #create the list of dicts
        for info in lst_shops:
            lst_info = info.replace("\n", "").split(";")
            dct_info = { lst_keys[i]: lst_info[i] for i in range(len(lst_info))}
            lst_csv_data.append(dct_info)

    return lst_csv_data

def convert_to_pdf(docx_name, path_pdf):
    """ converting the docx to pdf.
        @docx_name: the path to the docx file
        @path_pdf: path to the output file
    """
    convert(docx_name, path_pdf)

def create_pdfs(lst_data):
    """
        for each dict in the lst_data
        replace the placeholders and save as pdf
        @lst_data: list of dictionaries containing data from the csv
    """
    for info in lst_data:
        modify_docx(info)

    #clean up
    os.remove("tmp.docx")

def modify_docx(dct_info):
    """ create a new copy of the template and
        replace all the values.
    """
    document = Document("mail_template.docx")
    document.save("tmp.docx")
    document = Document("tmp.docx")

    for paragraph in document.paragraphs:
        inline = paragraph.runs
        for i, line in enumerate(inline):
            text = line.text
            #search the text and replace all placeholders
            for placeholder, value in dct_info.items():
                str_key = "<{}>".format(placeholder)
                if str_key in text:
                    text=text.replace(str_key, value)
            #save the updated text
            inline[i].text = text

    document.save("tmp.docx")

    #some small escaping done, to prevent the script from crashing
    output_name = dct_info.get("depot", "test").replace('"', "").replace(" ", "_")
    output_path = os.path.join("parsed", '{}.pdf'.format(output_name))

    convert_to_pdf("tmp.docx", output_path)

if __name__ == '__main__':
    #1. extract the keys and create a list of data.
    #2. Create parsed folder if it does not exists
    #3. Use the docx template, replace the data, save as pdf
    if not os.path.exists("./parsed"):
        os.mkdir("parsed")
    create_pdfs(get_info())
